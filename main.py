import os.path
import io

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
import os
from docx import Document
# If modifying these scopes, delete the file token.json.
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]


def download_files_from_drive():
    """Shows basic usage of the Drive v3 API.
    Prints the names and ids of the first 10 files the user has access to.
    """
    creds = None
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                "credentials.json", SCOPES
            )
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open("token.json", "w") as token:
            token.write(creds.to_json())

    try:
        service = build("drive", "v3", credentials=creds)

        # Call the Drive v3 API
        folder_id = "1X9DsQ814maw2RKfiT-lXF-Pn9YEBwf20"  # Replace with the actual folder ID
        results = service.files().list(
            q=f"'{folder_id}' in parents",
            pageSize=1000,
            fields="nextPageToken, files(id, name, mimeType)"
        ).execute()
        items = results.get("files", [])

        if not items:
            print("No files found in the specified folder.")
            return
        print("Files in the specified folder:")
        for item in items:
            if item['mimeType'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                print(f"{item['name']} ({item['id']})")
                file_id = item['id']
                request = service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                    print(f"Download {int(status.progress() * 100)}%.")
                
                # Save the file in the documents folder
                documents_folder = "documents"
                os.makedirs(documents_folder, exist_ok=True)
                file_path = os.path.join(documents_folder, item['name'])
                with open(file_path, "wb") as f:
                    f.write(fh.getvalue())
                print(f"Downloaded: {file_path}")
            else:
                print(f"Skipping non-DOCX file: {item['name']}")
    except HttpError as error:
        print(f"An error occurred: {error}")

def read_docx_from_documents():
    documents_folder = "documents"
    
    if not os.path.exists(documents_folder):
        print(f"The '{documents_folder}' folder does not exist.")
        return

    for filename in os.listdir(documents_folder):
        if filename.endswith(".docx"):
            file_path = os.path.join(documents_folder, filename)
            print(f"Reading file: {filename}")
            
            doc = Document(file_path)
            
            # for paragraph in doc.paragraphs:
            #     print(paragraph.text) 
            # Save the first paragraph to text.txt
            if doc.paragraphs:
                with open("text.txt", "w", encoding="utf-8") as f:
                    for i in range(min(20, len(doc.paragraphs))):
                        f.write(doc.paragraphs[i].text + "\n")
                print(f"First 10 paragraphs (or all if less than 10) saved to text.txt")
            else:
                print("No paragraphs found in the document.")

            
            
            print("\n" + "="*50 + "\n")  # Separator after the document
            return  # Exit after reading the first file


if __name__ == "__main__":
    # download_files_from_drive()
    read_docx_from_documents()